"""
FastAPI приложение для обработки файлов через Perplexity API.
"""

import uuid
import re
import json
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse

from wpd.init_core import init_core
from wpd.merge_with_docx import generate_docx_from_template
from wpd.fill_tables import fill_one_table_from_perplexity, _to_zero_based_table_index
from wpd.fill_result_table import fill_table_row_major
from wpd.tables_config import TABLE_SPECS, TABLE_INDEX_OFFSET
from wpd.table_prompts import TABLE_PROMPTS
from wpd.request_api import read_file_content, _load_chat_messages, _save_chat_messages
from dotenv import load_dotenv
from docx import Document

load_dotenv()
app = FastAPI(title="ГУАП - Формирование учебной программы")

# Папка для временных файлов и результатов
UPLOAD_DIR = Path("files/uploads")
RESULT_DIR = Path("files/results")
VARIABLES_DIR = Path("files/variables")  # Папка для сохранения JSON с переменными
TEMPLATE_PATH = Path("files/Шаблон.docx")

# Создаем папки если их нет
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
RESULT_DIR.mkdir(parents=True, exist_ok=True)
VARIABLES_DIR.mkdir(parents=True, exist_ok=True)


@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Главная страница с формой загрузки."""
    html_content = """
    <!DOCTYPE html>
    <html lang="ru">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>ГУАП - Формирование учебной программы</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                min-height: 100vh;
                display: flex;
                justify-content: center;
                align-items: center;
                padding: 20px;
            }
            .container {
                background: white;
                border-radius: 20px;
                box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
                padding: 40px;
                max-width: 800px;
                width: 100%;
            }
            h1 {
                color: #333;
                text-align: center;
                margin-bottom: 10px;
                font-size: 28px;
            }
            .subtitle {
                text-align: center;
                color: #666;
                margin-bottom: 30px;
                font-size: 16px;
            }
            .upload-area {
                border: 3px dashed #667eea;
                border-radius: 15px;
                padding: 40px;
                text-align: center;
                margin-bottom: 20px;
                transition: all 0.3s ease;
                cursor: pointer;
            }
            .upload-area:hover {
                background-color: #f8f9ff;
                border-color: #764ba2;
            }
            .upload-area.dragover {
                background-color: #e8ebff;
                border-color: #764ba2;
            }
            input[type="file"] {
                display: none;
            }
            .file-label {
                display: inline-block;
                padding: 12px 30px;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border-radius: 25px;
                cursor: pointer;
                font-size: 16px;
                font-weight: 600;
                transition: transform 0.2s ease;
            }
            .file-label:hover {
                transform: scale(1.05);
            }
            .file-name {
                margin-top: 15px;
                color: #333;
                font-weight: 500;
            }
            .file-name.empty {
                color: #999;
                font-style: italic;
            }
            button {
                width: 100%;
                padding: 15px;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                border: none;
                border-radius: 25px;
                font-size: 18px;
                font-weight: 600;
                cursor: pointer;
                transition: transform 0.2s ease, box-shadow 0.2s ease;
                margin-top: 20px;
            }
            button:hover:not(:disabled) {
                transform: translateY(-2px);
                box-shadow: 0 10px 20px rgba(102, 126, 234, 0.4);
            }
            button:disabled {
                opacity: 0.6;
                cursor: not-allowed;
            }
            .download-btn {
                display: none;
                margin-top: 20px;
                background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
            }
            .download-btn.show {
                display: block;
            }
            .status {
                margin-top: 20px;
                padding: 15px;
                border-radius: 10px;
                text-align: center;
                font-weight: 500;
                display: none;
            }
            .status.show {
                display: block;
            }
            .status.processing {
                background-color: #fff3cd;
                color: #856404;
            }
            .status.success {
                background-color: #d4edda;
                color: #155724;
            }
            .status.error {
                background-color: #f8d7da;
                color: #721c24;
            }
            .spinner {
                border: 3px solid #f3f3f3;
                border-top: 3px solid #667eea;
                border-radius: 50%;
                width: 30px;
                height: 30px;
                animation: spin 1s linear infinite;
                margin: 0 auto;
            }
            @keyframes spin {
                0% { transform: rotate(0deg); }
                100% { transform: rotate(360deg); }
            }
            .variables-section {
                margin-top: 30px;
                display: none;
            }
            .variables-section.show {
                display: block;
            }
            .variables-title {
                font-size: 20px;
                font-weight: 600;
                color: #333;
                margin-bottom: 20px;
            }
            .variable-item {
                display: flex;
                align-items: center;
                gap: 15px;
                margin-bottom: 15px;
                padding: 15px;
                background-color: #f8f9fa;
                border-radius: 10px;
                border: 1px solid #e0e0e0;
            }
            .variable-input {
                flex: 1;
                padding: 12px 15px;
                border: 2px solid #ddd;
                border-radius: 8px;
                font-size: 14px;
                transition: border-color 0.3s ease;
            }
            .variable-input:focus {
                outline: none;
                border-color: #667eea;
            }
            .variable-input:disabled {
                background-color: #e9ecef;
                cursor: not-allowed;
                color: #6c757d;
            }
            .variable-checkbox-container {
                display: flex;
                align-items: center;
                gap: 8px;
                min-width: 150px;
            }
            .variable-checkbox {
                width: 20px;
                height: 20px;
                cursor: pointer;
                accent-color: #667eea;
            }
            .variable-checkbox-label {
                font-size: 14px;
                color: #495057;
                cursor: pointer;
            }
            .tables-section {
                margin-top: 30px;
                display: none;
            }
            .tables-section.show {
                display: block;
            }
            .tables-title {
                font-size: 20px;
                font-weight: 600;
                color: #333;
                margin-bottom: 20px;
            }
            .table-item {
                display: flex;
                align-items: center;
                gap: 15px;
                margin-bottom: 15px;
                padding: 15px;
                background-color: #f8f9fa;
                border-radius: 10px;
                border: 1px solid #e0e0e0;
                cursor: pointer;
                transition: background-color 0.2s ease;
            }
            .table-item:hover {
                background-color: #e9ecef;
            }
            .table-name {
                flex: 1;
                font-weight: 500;
                color: #495057;
                font-size: 14px;
            }
            .table-name.empty {
                color: #999;
                font-style: italic;
            }
            .table-checkbox-container {
                display: flex;
                align-items: center;
                gap: 8px;
                min-width: 150px;
                flex-shrink: 0;
            }
            .table-checkbox {
                width: 20px;
                height: 20px;
                cursor: pointer;
                accent-color: #667eea;
            }
            .table-checkbox-label {
                font-size: 14px;
                color: #495057;
                cursor: pointer;
            }
            .modal {
                display: none;
                position: fixed;
                z-index: 1000;
                left: 0;
                top: 0;
                width: 100%;
                height: 100%;
                overflow: auto;
                background-color: rgba(0, 0, 0, 0.5);
            }
            .modal.show {
                display: flex;
                align-items: center;
                justify-content: center;
            }
            .modal-content {
                background-color: white;
                padding: 30px;
                border-radius: 15px;
                max-width: 90%;
                max-height: 90vh;
                overflow: auto;
                position: relative;
                box-shadow: 0 10px 40px rgba(0, 0, 0, 0.3);
            }
            .modal-header {
                display: flex;
                justify-content: space-between;
                align-items: center;
                margin-bottom: 20px;
            }
            .modal-title {
                font-size: 24px;
                font-weight: 600;
                color: #333;
            }
            .close-modal {
                font-size: 28px;
                font-weight: bold;
                color: #aaa;
                cursor: pointer;
                border: none;
                background: none;
                padding: 0;
                width: 30px;
                height: 30px;
                display: flex;
                align-items: center;
                justify-content: center;
            }
            .close-modal:hover {
                color: #000;
            }
            .table-editor {
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 20px;
            }
            .table-editor td, .table-editor th {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            .table-editor th {
                background-color: #667eea;
                color: white;
                font-weight: 600;
            }
            .table-editor td {
                background-color: white;
            }
            .table-cell-input {
                width: 100%;
                padding: 8px;
                border: none;
                font-size: 14px;
                background-color: transparent;
            }
            .table-cell-input:focus {
                outline: 2px solid #667eea;
                background-color: #f8f9ff;
            }
            .add-row-btn {
                padding: 10px 20px;
                background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
                color: white;
                border: none;
                border-radius: 8px;
                font-size: 14px;
                font-weight: 600;
                cursor: pointer;
                margin-top: 10px;
            }
            .add-row-btn:hover {
                transform: translateY(-2px);
                box-shadow: 0 5px 15px rgba(17, 153, 142, 0.4);
            }
            .add-row-btn:disabled {
                opacity: 0.5;
                cursor: not-allowed;
                transform: none;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>ГУАП</h1>
            <p class="subtitle">Формирование учебной программы</p>
            
            <div class="upload-area" id="uploadArea">
                <input type="file" id="fileInput" accept=".docx" />
                <label for="fileInput" class="file-label">Загрузить учебник</label>
                <div class="file-name empty" id="fileName">Файл не выбран</div>
            </div>
            
            <div class="variables-section" id="variablesSection">
                <div class="variables-title">Настройка переменных:</div>
                <div id="variablesList"></div>
            </div>
            
            <div class="tables-section" id="tablesSection">
                <div class="tables-title">Настройка таблиц:</div>
                <div id="tablesList"></div>
            </div>
            
            <button id="submitBtn" onclick="submitFile()" disabled>Отправить</button>
            
            <div class="status" id="status"></div>
            
            <a id="downloadBtn" class="download-btn" href="#" download style="display: none; text-decoration: none; color: white; text-align: center; padding: 15px; border-radius: 25px; font-size: 18px; font-weight: 600; margin-top: 20px;">
                Скачать результат
            </a>
        </div>
        
        <script>
            const fileInput = document.getElementById('fileInput');
            const fileName = document.getElementById('fileName');
            const submitBtn = document.getElementById('submitBtn');
            const uploadArea = document.getElementById('uploadArea');
            const status = document.getElementById('status');
            const downloadBtn = document.getElementById('downloadBtn');
            const variablesSection = document.getElementById('variablesSection');
            const variablesList = document.getElementById('variablesList');
            const tablesSection = document.getElementById('tablesSection');
            const tablesList = document.getElementById('tablesList');
            
            let variablesData = [];
            let tablesData = [];
            let currentEditingTableIndex = null;
            
            // Обработка выбора файла
            fileInput.addEventListener('change', async function(e) {
                const file = e.target.files[0];
                if (file) {
                    fileName.textContent = file.name;
                    fileName.classList.remove('empty');
                    submitBtn.disabled = true; // Отключаем кнопку до загрузки переменных
                    
                    // Загружаем переменные и таблицы с сервера
                    try {
                        // Загружаем переменные
                        const varsResponse = await fetch('/template/variables');
                        if (!varsResponse.ok) {
                            throw new Error('Не удалось загрузить переменные');
                        }
                        const varsData = await varsResponse.json();
                        variablesData = varsData.variables;
                        renderVariables();
                        variablesSection.classList.add('show');
                        
                        // Загружаем таблицы
                        const tablesResponse = await fetch('/template/tables');
                        if (!tablesResponse.ok) {
                            throw new Error('Не удалось загрузить таблицы');
                        }
                        const tablesResult = await tablesResponse.json();
                        tablesData = tablesResult.tables;
                        renderTables();
                        tablesSection.classList.add('show');
                    } catch (error) {
                        status.className = 'status error show';
                        status.textContent = 'Ошибка загрузки данных: ' + error.message;
                    }
                } else {
                    fileName.textContent = 'Файл не выбран';
                    fileName.classList.add('empty');
                    submitBtn.disabled = true;
                    variablesSection.classList.remove('show');
                    tablesSection.classList.remove('show');
                    variablesData = [];
                    tablesData = [];
                }
            });
            
            // Отрисовка переменных
            function renderVariables() {
                variablesList.innerHTML = '';
                
                variablesData.forEach((variable, index) => {
                    const variableItem = document.createElement('div');
                    variableItem.className = 'variable-item';
                    
                    const input = document.createElement('input');
                    input.type = 'text';
                    input.className = 'variable-input';
                    input.placeholder = variable.name;  // Название переменной как placeholder
                    input.value = variable.value || '';  // Только введенное значение
                    input.disabled = variable.auto_generate;
                    input.id = 'input-' + index;
                    
                    input.addEventListener('input', function() {
                        if (this.disabled) return;
                        // Сохраняем введенное значение
                        variablesData[index].value = this.value;
                    });
                    
                    const checkboxContainer = document.createElement('div');
                    checkboxContainer.className = 'variable-checkbox-container';
                    
                    const checkbox = document.createElement('input');
                    checkbox.type = 'checkbox';
                    checkbox.className = 'variable-checkbox';
                    checkbox.checked = variable.auto_generate;
                    checkbox.addEventListener('change', function() {
                        variablesData[index].auto_generate = this.checked;
                        input.disabled = this.checked;
                        if (this.checked) {
                            // Если включена автогенерация, очищаем поле
                            input.value = '';
                            variablesData[index].value = '';
                        }
                    });
                    
                    const checkboxLabel = document.createElement('label');
                    checkboxLabel.className = 'variable-checkbox-label';
                    checkboxLabel.textContent = 'Автогенерация';
                    checkboxLabel.htmlFor = 'var-' + index;
                    checkbox.id = 'var-' + index;
                    
                    checkboxContainer.appendChild(checkbox);
                    checkboxContainer.appendChild(checkboxLabel);
                    
                    variableItem.appendChild(input);
                    variableItem.appendChild(checkboxContainer);
                    
                    variablesList.appendChild(variableItem);
                });
                
                // Включаем кнопку отправки после отрисовки переменных
                submitBtn.disabled = false;
            }
            
            // Отрисовка таблиц
            function renderTables() {
                tablesList.innerHTML = '';
                
                tablesData.forEach((table, index) => {
                    const tableItem = document.createElement('div');
                    tableItem.className = 'table-item';
                    tableItem.onclick = function(e) {
                        // Не открываем редактор при клике на чекбокс
                        if (e.target.type === 'checkbox' || e.target.tagName === 'LABEL') {
                            return;
                        }
                        openTableEditor(index);
                    };
                    
                    const tableName = document.createElement('div');
                    // Показываем название из поля name, если оно заполнено, иначе номер таблицы
                    const displayName = table.name && table.name.trim() ? table.name : `Таблица #${table.table_number}`;
                    tableName.className = 'table-name' + (table.name && table.name.trim() ? '' : ' empty');
                    tableName.textContent = displayName;
                    
                    const checkboxContainer = document.createElement('div');
                    checkboxContainer.className = 'table-checkbox-container';
                    
                    const checkbox = document.createElement('input');
                    checkbox.type = 'checkbox';
                    checkbox.className = 'table-checkbox';
                    checkbox.checked = table.should_fill_with_ai || false;
                    checkbox.onclick = function(e) {
                        e.stopPropagation();
                    };
                    checkbox.addEventListener('change', function() {
                        tablesData[index].should_fill_with_ai = this.checked;
                    });
                    
                    const checkboxLabel = document.createElement('label');
                    checkboxLabel.className = 'table-checkbox-label';
                    checkboxLabel.textContent = 'Автогенерация';
                    checkboxLabel.htmlFor = 'table-' + index;
                    checkbox.id = 'table-' + index;
                    checkboxLabel.onclick = function(e) {
                        e.stopPropagation();
                    };
                    
                    checkboxContainer.appendChild(checkbox);
                    checkboxContainer.appendChild(checkboxLabel);
                    
                    tableItem.appendChild(tableName);
                    tableItem.appendChild(checkboxContainer);
                    
                    tablesList.appendChild(tableItem);
                });
            }
            
            // Открытие модального окна для редактирования таблицы
            function openTableEditor(tableIndex) {
                currentEditingTableIndex = tableIndex;
                const table = tablesData[tableIndex];
                
                // Создаем модальное окно, если его еще нет
                let modal = document.getElementById('tableModal');
                if (!modal) {
                    modal = document.createElement('div');
                    modal.id = 'tableModal';
                    modal.className = 'modal';
                    document.body.appendChild(modal);
                }
                
                modal.innerHTML = `
                    <div class="modal-content">
                        <div class="modal-header">
                            <h2 class="modal-title">${table.name || `Таблица #${table.table_number}`}</h2>
                            <button class="close-modal" onclick="closeTableEditor()">&times;</button>
                        </div>
                        <div id="tableEditorContent"></div>
                        ${table.can_add_rows ? '<button class="add-row-btn" onclick="addTableRow()">Добавить строку</button>' : ''}
                    </div>
                `;
                
                // Отрисовываем таблицу для редактирования
                const editorContent = document.getElementById('tableEditorContent');
                editorContent.innerHTML = '';
                const tableElement = document.createElement('table');
                tableElement.className = 'table-editor';
                
                // Отрисовываем строки таблицы
                table.data.forEach((row, rowIndex) => {
                    const tr = document.createElement('tr');
                    row.forEach((cell, cellIndex) => {
                        const td = document.createElement('td');
                        const input = document.createElement('input');
                        input.type = 'text';
                        input.className = 'table-cell-input';
                        input.value = cell;
                        input.dataset.rowIndex = rowIndex;
                        input.dataset.cellIndex = cellIndex;
                        input.addEventListener('input', function() {
                            tablesData[tableIndex].data[rowIndex][cellIndex] = this.value;
                        });
                        td.appendChild(input);
                        tr.appendChild(td);
                    });
                    tableElement.appendChild(tr);
                });
                
                editorContent.appendChild(tableElement);
                modal.classList.add('show');
                
                // Закрытие по клику вне модального окна
                modal.onclick = function(e) {
                    if (e.target === modal) {
                        closeTableEditor();
                    }
                };
            }
            
            // Закрытие модального окна
            function closeTableEditor() {
                const modal = document.getElementById('tableModal');
                if (modal) {
                    modal.classList.remove('show');
                    currentEditingTableIndex = null;
                }
            }
            
            // Добавление строки в таблицу
            function addTableRow() {
                if (currentEditingTableIndex === null) return;
                
                const table = tablesData[currentEditingTableIndex];
                if (!table.can_add_rows) return;
                
                // Создаем новую строку с пустыми ячейками
                const numCols = table.num_cols || (table.data.length > 0 ? table.data[0].length : 0);
                const newRow = Array(numCols).fill('');
                table.data.push(newRow);
                table.num_rows = table.data.length;
                
                // Перерисовываем таблицу
                openTableEditor(currentEditingTableIndex);
            }
            
            // Делаем функции доступными глобально для onclick
            window.openTableEditor = openTableEditor;
            window.closeTableEditor = closeTableEditor;
            window.addTableRow = addTableRow;
            
            // Drag and drop
            uploadArea.addEventListener('dragover', function(e) {
                e.preventDefault();
                uploadArea.classList.add('dragover');
            });
            
            uploadArea.addEventListener('dragleave', function(e) {
                e.preventDefault();
                uploadArea.classList.remove('dragover');
            });
            
            uploadArea.addEventListener('drop', function(e) {
                e.preventDefault();
                uploadArea.classList.remove('dragover');
                const files = e.dataTransfer.files;
                if (files.length > 0) {
                    fileInput.files = files;
                    fileInput.dispatchEvent(new Event('change'));
                }
            });
            
            // Отправка файла
            async function submitFile() {
                const file = fileInput.files[0];
                if (!file) {
                    alert('Пожалуйста, выберите файл');
                    return;
                }
                
                if (variablesData.length === 0) {
                    alert('Загрузите переменные');
                    return;
                }
                
                // Подготавливаем данные для отправки
                const variablesToSend = variablesData.map(variable => ({
                    name: variable.name,
                    value: variable.value || '',
                    auto_generate: variable.auto_generate || false
                }));
                
                const tablesToSend = tablesData.map(table => ({
                    table_index: table.table_index,
                    table_number: table.table_number,
                    name: table.name || '',
                    num_rows: table.num_rows,
                    num_cols: table.num_cols,
                    data: table.data,
                    can_add_rows: table.can_add_rows || false,
                    should_fill_with_ai: table.should_fill_with_ai || false
                }));
                
                const formData = new FormData();
                formData.append('file', file);
                formData.append('variables', JSON.stringify({
                    variables: variablesToSend,
                    count: variablesToSend.length
                }));
                formData.append('tables', JSON.stringify({
                    tables: tablesToSend,
                    count: tablesToSend.length
                }));
                
                submitBtn.disabled = true;
                status.className = 'status processing show';
                status.innerHTML = '<div class="spinner"></div><br>Обработка файла...';
                downloadBtn.style.display = 'none';
                
                try {
                    const response = await fetch('/upload', {
                        method: 'POST',
                        body: formData
                    });
                    
                    if (!response.ok) {
                        const error = await response.json();
                        throw new Error(error.detail || 'Ошибка при обработке файла');
                    }
                    
                    const result = await response.json();
                    
                    status.className = 'status success show';
                    status.textContent = 'Файл успешно обработан!';
                    
                    downloadBtn.href = `/download/${result.file_id}`;
                    downloadBtn.style.display = 'block';
                    
                } catch (error) {
                    status.className = 'status error show';
                    status.textContent = 'Ошибка: ' + error.message;
                } finally {
                    submitBtn.disabled = false;
                }
            }
        </script>
    </body>
    </html>
    """
    return html_content


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    variables: str = Form(...),
    tables: str = Form(None)
):
    """
    Загружает файл от пользователя, JSON с переменными и таблицами, обрабатывает через ядро.
    
    Args:
        file: загруженный файл учебника
        variables: JSON строка с переменными в формате {"variables": [...], "count": N}
        tables: JSON строка с таблицами в формате {"tables": [...], "count": N} (опционально)
    
    Returns:
        dict с file_id для скачивания результата
    """
    # Проверяем расширение файла
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Поддерживаются только файлы .docx")
    
    # Генерируем уникальный ID для сессии
    file_id = str(uuid.uuid4())
    
    # Сохраняем загруженный файл
    uploaded_file_path = UPLOAD_DIR / f"{file_id}_{file.filename}"
    with open(uploaded_file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # Парсим и сохраняем JSON с переменными
    try:
        variables_data = json.loads(variables)
        variables_file_path = VARIABLES_DIR / f"{file_id}_variables.json"
        with open(variables_file_path, 'w', encoding='utf-8') as f:
            json.dump(variables_data, f, ensure_ascii=False, indent=2)
        print(f"Сохранен JSON с переменными: {variables_file_path}")
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Ошибка парсинга JSON с переменными: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка сохранения переменных: {str(e)}")
    
    # Парсим и сохраняем JSON с таблицами (если передан)
    if tables:
        try:
            tables_data = json.loads(tables)
            tables_file_path = VARIABLES_DIR / f"{file_id}_tables.json"
            with open(tables_file_path, 'w', encoding='utf-8') as f:
                json.dump(tables_data, f, ensure_ascii=False, indent=2)
            print(f"Сохранен JSON с таблицами: {tables_file_path}")
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"Ошибка парсинга JSON с таблицами: {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Ошибка сохранения таблиц: {str(e)}")
    
    try:
        # Путь к шаблону (фиксированный)
        template_path = str(TEMPLATE_PATH)
        if not Path(template_path).exists():
            raise HTTPException(
                status_code=500,
                detail=f"Шаблон не найден: {template_path}"
            )
        
        # Путь к результату
        result_path = RESULT_DIR / f"{file_id}_result.docx"
        
        # Шаг 1: Обработка переменных из JSON
        variables_list = variables_data.get('variables', [])
        
        # Создаем словарь переменных ТОЛЬКО из JSON (не из шаблона!)
        # Это нужно для условных блоков - переменные должны быть в контексте, даже если пустые
        all_variables_dict = {}
        
        # Добавляем все переменные из JSON в словарь
        # Это нужно для поддержки условных блоков и новых переменных
        for var in variables_list:
            name = var.get('name', '').strip()
            if not name:
                continue
            
            # Инициализируем переменную пустой строкой (для условных блоков)
            if name not in all_variables_dict:
                all_variables_dict[name] = ""
            
            if not var.get('auto_generate', False):
                # Переменные с auto_generate = false - берем значение из JSON
                value = var.get('value', '').strip()
                # Заменяем символ ; на ; + символ новой строки (для создания элементов списка)
                value = value.replace(';', ';\n')
                if value:  # Обновляем только если значение не пустое
                    all_variables_dict[name] = value
        
        # Шаг 3: Обработка переменных с auto_generate = true через ИИ
        auto_generate_variables = [var for var in variables_list if var.get('auto_generate', False)]
        # Создаем множество имен переменных, которые запрошены для генерации
        auto_generate_names = {var.get('name', '').strip() for var in auto_generate_variables if var.get('name', '').strip()}
        thread_id = None
        
        if auto_generate_variables:
            print(f"Обрабатываем {len(auto_generate_variables)} переменных через ИИ...")
            prompt = (
            "Привет, ты профессиональный эксперт-методист с 15-летним стажем работы в сфере. "
            "Я прикрепляю для тебя 2 файла: шаблон Рабочей программы дисциплины от ВУЗа, а также учебные материалы. "
            "Тебе нужно заполнить шаблон Рабочей программы дисциплины, основываясь на учебных материалах, "
            "которые содержат всё, что планируется реализовать в программе на семестр. "
            "Для того, чтобы это сделать, для начала тебе нужно проанализировать шаблон и то, чего там не хватает "
            "(что нужно заполнить) (все эти места являются как бы переменными и отмечены двойными фигурными скобками, "
            "внутри них содержится краткое описание того, что там должно быть), а затем, проанализировав учебные материалы, "
            "найти те недостающие 'переменные', которые нужно заполнить в шаблоне. "
            "Ты должен выбрать и вернуть мне именно то, что непосредственно прямо указано в материалах. "
            "Те переменные, которые там не упоминаются, или которые ты не смог найти - просто пропускай и не вноси в финальный результат, "
            "который ты будешь возвращать мне. "
            "Возвращать данные мне ты должен в формате ключ:значение; ключ:значение;..., "
            "где ключ - это полное название переменной, как в шаблоне, а значение - то значение, которое ты для нее нашел. "
            "Не добавляй в ответ никакие специальные символы, разделения строк и так далее. "
            "Когда выводишь список переменных и их значений не оборачивай ключи или значения в спец символы. "
                "Символ новой строки после знака точки с запятой тоже ставить не нужно"
            )
            
            # Вызываем API для получения значений переменных
            from wpd.request_api import call_api_in_one
            answer, thread_id = call_api_in_one(
                file1_path=str(template_path),  # Используем оригинальный шаблон
                file2_path=str(uploaded_file_path),  # Загруженный учебник от пользователя
                prompt=prompt,
                model="sonar"
            )
            
            # Парсим ответ от API и обновляем значения в словаре
            from wpd.merge_with_docx import _parse_pairs_from_text
            ai_variables = _parse_pairs_from_text(answer)
            
            # Обновляем значения переменных от ИИ ТОЛЬКО для тех, которые были запрошены для генерации
            # и не перезаписываем уже заполненные вручную
            updated_count = 0
            for key, value in ai_variables:
                # Проверяем, что переменная была запрошена для генерации
                if key in auto_generate_names and key in all_variables_dict:
                    # Обновляем только если переменная еще не заполнена вручную
                    if not all_variables_dict[key]:
                        # Заменяем символ ; на ; + символ новой строки (для создания элементов списка)
                        value = value.replace(';', ';\n')
                        all_variables_dict[key] = value
                        updated_count += 1
            
            print(f"Переменные с автогенерацией заполнены: {updated_count} из {len(ai_variables)} полученных от ИИ")
        
        # Шаг 4: Генерируем документ со всеми переменными (включая пустые для условных блоков)
        print(f"Генерируем документ с {len(all_variables_dict)} переменными...")
        generate_docx_from_template(
            {},  # Пустой словарь, так как все переменные уже в all_variables_dict
            template_path, 
            str(result_path),
            all_variables=all_variables_dict  # Передаем все переменные для поддержки условных блоков
        )
        print(f"Создан файл с переменными: {result_path}")
        
        # Шаг 3: Обработка таблиц из JSON
        if tables:
            tables_list = tables_data.get('tables', [])
            
            # Если thread_id еще не создан (не было переменных с автогенерацией), создаем его
            if not thread_id:
                import uuid as _uuid
                thread_id = str(_uuid.uuid4())
                
                # Загружаем файлы в историю чата для контекста
                messages = _load_chat_messages(thread_id)
                if not messages:
                    file1_content = read_file_content(str(template_path))
                    file2_content = read_file_content(str(uploaded_file_path))
                    messages = [
                        {"role": "system", "content": "Вы — полезный ассистент, который анализирует файлы и отвечает на вопросы."},
                        {
                            "role": "user",
                            "content": (
                                f"Файл 1 ({Path(template_path).name}):\n{file1_content}\n\n"
                                f"Файл 2 ({Path(uploaded_file_path).name}):\n{file2_content}"
                            ),
                        }
                    ]
                    _save_chat_messages(thread_id, messages)
            
            for table in tables_list:
                table_index = table.get('table_index')
                should_fill_with_ai = table.get('should_fill_with_ai', False)
                
                # Находим соответствующую TableFillSpec по table_index
                spec = None
                for s in TABLE_SPECS:
                    if s.table_index == table_index:
                        spec = s
                        break
                
                if not spec:
                    print(f"Предупреждение: Не найдена конфигурация для таблицы с table_index={table_index}")
                    continue
                
                if not should_fill_with_ai:
                    # Заполняем таблицу из JSON данных
                    print(f"Заполняем таблицу {table_index} из JSON данных...")
                    table_data = table.get('data', [])
                    # Преобразуем двумерный массив в плоский список (row-major order)
                    # Пропускаем строки до start_row (обычно это заголовки)
                    # И колонки до start_col в каждой строке
                    flat_values = []
                    for row_idx, row in enumerate(table_data):
                        if row_idx >= spec.start_row:  # Пропускаем заголовки
                            # Пропускаем колонки до start_col и берем только нужные ячейки
                            row_cells = [str(cell) for cell in row[spec.start_col:]]
                            flat_values.extend(row_cells)
                    print(f"Извлечено {len(flat_values)} значений из JSON данных (строк: {len(table_data)}, start_row: {spec.start_row}, start_col: {spec.start_col})")
                    
                    # Преобразуем table_index в индекс для python-docx
                    doc_table_index = _to_zero_based_table_index(
                        spec.table_index,
                        index_base=1,
                        table_index_offset=TABLE_INDEX_OFFSET
                    )
                    
                    fill_table_row_major(
                        result_docx_path=str(result_path),
                        values=flat_values,
                        table_index=doc_table_index,
                        cols_per_row=spec.cols_per_row,
                        start_row=spec.start_row,
                        start_col=spec.start_col,
                    )
                    print(f"Таблица {table_index} заполнена из JSON")
                else:
                    # Заполняем таблицу через ИИ
                    print(f"Заполняем таблицу {table_index} через ИИ...")
                    prompt_idx = spec.prompt_idx
                    prompt = TABLE_PROMPTS[prompt_idx]
                    
                    fill_one_table_from_perplexity(
                        result_docx_path=str(result_path),
                        table_index=spec.table_index,
                        cols_per_row=spec.cols_per_row,
                        start_row=spec.start_row,
                        start_col=spec.start_col,
            prompt=prompt,
            model="sonar",
                        thread_id=thread_id,
                        index_base=1,
                        table_index_offset=TABLE_INDEX_OFFSET,
        )
                    print(f"Таблица {table_index} заполнена через ИИ")
        
        return {"file_id": file_id, "message": "Файл успешно обработан"}
        
    except ValueError as e:
        # Ошибки валидации (например, отсутствие API ключа)
        error_msg = str(e)
        if uploaded_file_path.exists():
            uploaded_file_path.unlink()
        import traceback
        print(f"Ошибка валидации: {error_msg}")
        traceback.print_exc()
        # Более понятное сообщение для пользователя
        if "PPLX_API_KEY" in error_msg:
            raise HTTPException(
                status_code=500, 
                detail="Ошибка конфигурации: API ключ не установлен. Проверьте переменную окружения PPLX_API_KEY в .env файле."
            )
        raise HTTPException(status_code=500, detail=f"Ошибка конфигурации: {error_msg}")
    except FileNotFoundError as e:
        # Файл не найден
        error_msg = str(e)
        if uploaded_file_path.exists():
            uploaded_file_path.unlink()
        print(f"Файл не найден: {error_msg}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500, 
            detail=f"Файл не найден: {error_msg}. Убедитесь что файл шаблона files/Шаблон.docx существует."
        )
    except Exception as e:
        # Общие ошибки
        error_msg = str(e)
        if uploaded_file_path.exists():
            uploaded_file_path.unlink()
        import traceback
        print(f"Ошибка при обработке: {error_msg}")
        traceback.print_exc()
        # Более детальное сообщение об ошибке
        if "PPLX_API_KEY" in error_msg or "API" in error_msg:
            raise HTTPException(
                status_code=500, 
                detail=f"Ошибка API: Проверьте что PPLX_API_KEY установлен в переменных окружения. Детали: {error_msg}"
            )
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке: {error_msg}")


@app.get("/download/{file_id}")
async def download_file(file_id: str):
    """
    Скачивает обработанный файл result.docx.
    
    Args:
        file_id: ID файла из ответа /upload
        
    Returns:
        FileResponse с файлом result.docx
    """
    result_file = RESULT_DIR / f"{file_id}_result.docx"
    
    if not result_file.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")
    
    return FileResponse(
        path=str(result_file),
        filename="result.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/template/variables")
async def get_template_variables():
    """
    Возвращает список всех переменных из шаблона.
    Переменные имеют формат {{переменная}}.
    Результат в формате JSON для использования на frontend.
    """
    try:
        if not TEMPLATE_PATH.exists():
            raise HTTPException(
                status_code=404,
                detail=f"Шаблон не найден: {TEMPLATE_PATH}"
            )
        
        # Извлекаем переменные из шаблона
        doc = Document(str(TEMPLATE_PATH))
        variables = set()
        
        # Регулярное выражение для поиска переменных в формате {{переменная}}
        pattern = re.compile(r'\{\{\s*([^}]+)\s*\}\}')
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = pattern.findall(text)
            for match in matches:
                var_name = match.strip()
                if var_name:
                    variables.add(var_name)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    matches = pattern.findall(text)
                    for match in matches:
                        var_name = match.strip()
                        if var_name:
                            variables.add(var_name)
        
        # Формируем список переменных
        # Каждая переменная содержит:
        # - name: название переменной
        # - value: значение (пустое по умолчанию)
        # - auto_generate: флаг нужно ли пытаться сгенерировать переменную на основе учебного материала (по умолчанию false)
        variables_list = [
            {
                "name": var_name,
                "value": "",
                "auto_generate": False  # По умолчанию переменные не будут генерироваться автоматически
            }
            for var_name in sorted(variables)
        ]
        
        return JSONResponse(content={
            "variables": variables_list,
            "count": len(variables_list)
        })
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при извлечении переменных: {str(e)}"
        )


@app.get("/template/tables")
async def get_template_tables():
    """
    Возвращает список таблиц из JSON файла template_tables.json.
    Если JSON файла нет, выбрасывает ошибку 404.
    """
    try:
        # Проверяем оба возможных расположения JSON файла
        json_file_paths = [
            Path("template_tables.json"),  # В корне проекта
            Path("files/template_tables.json")  # В папке files
        ]
        
        json_file_path = None
        for path in json_file_paths:
            if path.exists():
                json_file_path = path
                break
        
        # Если JSON файл не найден, выбрасываем ошибку
        if not json_file_path:
            raise HTTPException(
                status_code=404,
                detail="Файл template_tables.json не найден. Создайте его с помощью скрипта extract_template_tables.py"
            )
        
        # Загружаем и возвращаем данные из JSON файла
        try:
            with open(json_file_path, 'r', encoding='utf-8') as f:
                saved_data = json.load(f)
                if 'tables' in saved_data:
                    print(f"Загружены таблицы из JSON: {json_file_path} (количество: {len(saved_data['tables'])})")
                    # Возвращаем таблицы из JSON файла как есть
                    return JSONResponse(content={
                        "tables": saved_data['tables'],
                        "count": len(saved_data['tables'])
                    })
                else:
                    raise HTTPException(
                        status_code=500,
                        detail=f"В файле {json_file_path} отсутствует поле 'tables'"
                    )
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка парсинга JSON файла {json_file_path}: {str(e)}"
            )
        except HTTPException:
            raise  # Пробрасываем HTTPException без изменений
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка при загрузке файла {json_file_path}: {str(e)}"
            )
    except HTTPException:
        raise  # Пробрасываем HTTPException без изменений
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при загрузке таблиц: {str(e)}"
        )


@app.get("/favicon.ico")
async def favicon():
    """Обработка favicon чтобы избежать 404 ошибок"""
    from fastapi.responses import Response
    return Response(status_code=204)  # No Content


@app.on_event("startup")
async def startup_event():
    """Событие запуска приложения"""
    import os
    print("=" * 60)
    print("🚀 Веб-сервер запущен")
    print("=" * 60)
    print(f"📍 URL: http://0.0.0.0:8000")
    
    # Проверка переменных окружения
    pplx_key = os.getenv("PPLX_API_KEY")
    if pplx_key:
        print(f"✅ PPLX_API_KEY установлен (длина: {len(pplx_key)} символов)")
    else:
        print("⚠️  PPLX_API_KEY НЕ установлен! API запросы не будут работать.")
    
    telegram_token = os.getenv("TELEGRAM_BOT_TOKEN")
    if telegram_token:
        print(f"✅ TELEGRAM_BOT_TOKEN установлен")
    else:
        print("ℹ️  TELEGRAM_BOT_TOKEN не установлен (Telegram бот не будет работать)")
    
    # Проверка шаблона
    if TEMPLATE_PATH.exists():
        print(f"✅ Шаблон найден: {TEMPLATE_PATH}")
    else:
        print(f"❌ Шаблон НЕ найден: {TEMPLATE_PATH}")
    
    print("=" * 60)


if __name__ == "__main__":
    import uvicorn
    print("Запуск веб-сервера...")
    print("Для запуска с Telegram ботом используйте: python run_all.py")
    uvicorn.run(app, host="0.0.0.0", port=8000)

