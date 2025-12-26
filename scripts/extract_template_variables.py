"""
Скрипт для извлечения переменных из шаблона docx файла.
Переменные имеют формат {{переменная}}
Результат сохраняется в JSON файл для отправки на frontend.
"""

import re
import json
from pathlib import Path
from docx import Document

def extract_variables_from_docx(template_path: str) -> list[dict]:
    """
    Извлекает все переменные из docx файла.
    Переменные имеют формат {{переменная}}
    
    Args:
        template_path: путь к файлу шаблона
        
    Returns:
        список словарей с информацией о переменных
    """
    try:
        doc = Document(template_path)
        variables = set()  # Используем set для уникальности
        
        # Регулярное выражение для поиска переменных в формате {{переменная}}
        # Поддерживает пробелы внутри скобок: {{ переменная }} или {{переменная}}
        pattern = re.compile(r'\{\{\s*([^}]+)\s*\}\}')
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            text = paragraph.text
            matches = pattern.findall(text)
            for match in matches:
                var_name = match.strip()
                if var_name:
                    variables.add(var_name)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    matches = pattern.findall(text)
                    for match in matches:
                        var_name = match.strip()
                        if var_name:
                            variables.add(var_name)
        
        # Преобразуем в список словарей для JSON
        # Каждая переменная будет объектом с полями:
        # - name: название переменной
        # - value: значение (пустое по умолчанию)
        # - auto_generate: флаг нужно ли пытаться сгенерировать переменную на основе учебного материала (по умолчанию false)
        variables_list = [
            {
                "name": var_name,
                "value": "",
                "auto_generate": False  # По умолчанию переменные не будут генерироваться автоматически
            }
            for var_name in sorted(variables)  # Сортируем для удобства
        ]
        
        return variables_list
        
    except Exception as e:
        raise ValueError(f"Ошибка при чтении файла {template_path}: {e}")


def main():
    """Главная функция для извлечения переменных и сохранения в JSON"""
    # Путь к шаблону (относительно корня проекта)
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    template_path = project_root / "files" / "Шаблон.docx"
    
    if not template_path.exists():
        print(f"Ошибка: файл {template_path} не найден!")
        return
    
    print(f"Чтение файла: {template_path}")
    
    # Извлекаем переменные
    variables = extract_variables_from_docx(str(template_path))
    
    print(f"Найдено переменных: {len(variables)}")
    
    # Формируем JSON структуру
    result = {
        "variables": variables,
        "count": len(variables)
    }
    
    # Сохраняем в JSON файл (в корне проекта)
    output_path = project_root / "template_variables.json"
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    print(f"Результат сохранен в: {output_path}")
    print("\nСписок найденных переменных:")
    for i, var in enumerate(variables, 1):
        print(f"  {i}. {var['name']}")


if __name__ == "__main__":
    main()

