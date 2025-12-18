from __future__ import annotations

import json
import re
from typing import List, Optional, Sequence

from docx import Document

from wpd.request_api import client, read_file_content, DEFAULT_CHAT_STORE, _load_chat_messages, _save_chat_messages
import os
from pathlib import Path


def find_table_index_by_headers(
    result_docx_path: str,
    header_substrings: Sequence[str],
    header_row: int = 0,
) -> int:
    """
    Ищет таблицу в docx по подстрокам в строке заголовка.
    Возвращает индекс таблицы (как в python-docx: doc.tables[index]).
    """
    doc = Document(result_docx_path)
    need = [s.casefold() for s in header_substrings]

    for i, table in enumerate(doc.tables):
        if len(table.rows) <= header_row:
            continue
        header_text = " | ".join(c.text.strip() for c in table.rows[header_row].cells).casefold()
        if all(s in header_text for s in need):
            return i

    raise ValueError(
        "Не удалось найти таблицу по заголовкам. "
        f"Искали подстроки: {list(header_substrings)}. "
        "Укажите table_index вручную."
    )


def _extract_list_values(text: str) -> List[str]:
    """
    Пытается достать список значений из ответа модели.

    Поддерживаемые форматы:
    - JSON массив: ["a", "b", "c"]
    - Маркированный список:
      - a
      - b
    - Нумерованный список:
      1) a
      2) b
    - Одна строка через ';' или ','
    """
    if not text:
        return []

    s = text.strip()

    # 1) JSON целиком
    try:
        obj = json.loads(s)
        if isinstance(obj, list):
            return [str(x).strip() for x in obj if str(x).strip()]
    except Exception:
        pass

    # 2) JSON-массив как подстрока
    m = re.search(r"\[[\s\S]*\]", s)
    if m:
        try:
            obj = json.loads(m.group(0))
            if isinstance(obj, list):
                return [str(x).strip() for x in obj if str(x).strip()]
        except Exception:
            pass

    # 3) Маркеры/нумерация по строкам
    values: List[str] = []
    for line in s.replace("\r\n", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        line = re.sub(r"^[-•]\s*", "", line)            # - item / • item
        line = re.sub(r"^\d+[\.\)]\s*", "", line)       # 1. item / 1) item
        if line:
            values.append(line)
    if len(values) >= 2:
        return values

    # 4) Разделители
    if ";" in s:
        return [p.strip() for p in s.split(";") if p.strip()]
    if "," in s:
        return [p.strip() for p in s.split(",") if p.strip()]

    return [s] if s else []


def fill_table_row_major(
    result_docx_path: str,
    values: Sequence[str],
    table_index: int = 1,
    cols_per_row: int = 3,
    start_row: int = 0,
    start_col: int = 0,
) -> str:
    """
    Заполняет таблицу в docx по индексу `table_index` значениями из списка
    слева-направо, сверху-вниз (row-major order).

    Логика заполнения:
    - ширина "полезной" области задаётся `cols_per_row`
    - values[0] -> cell[start_row + 0][start_col + 0]
    - values[1] -> cell[start_row + 0][start_col + 1]
    - values[2] -> cell[start_row + 0][start_col + 2]
    - values[3] -> cell[start_row + 1][start_col + 0]
    - и т.д.
    - если строк не хватает — добавляем строки
    """
    doc = Document(result_docx_path)
    if table_index >= len(doc.tables):
        raise ValueError(
            f"В документе {result_docx_path} нет таблицы с индексом {table_index}. "
            f"Доступно таблиц: {len(doc.tables)}"
        )

    table = doc.tables[table_index]
    if cols_per_row <= 0:
        raise ValueError("cols_per_row должен быть > 0")
    if start_row < 0 or start_col < 0:
        raise ValueError("start_row и start_col должны быть >= 0")
    if len(table.rows) == 0:
        raise ValueError("Таблица пуста (нет строк).")
    
    # Используем максимальное количество ячеек в строке вместо len(table.columns)
    # для корректной работы с объединенными ячейками
    max_row_cells = max((len(r.cells) for r in table.rows), default=0)
    if max_row_cells == 0:
        raise ValueError("Таблица пуста (нет ячеек).")
    if start_col + cols_per_row > max_row_cells:
        raise ValueError(
            f"Недостаточно колонок: нужно минимум {start_col + cols_per_row}, "
            f"а в таблице максимум {max_row_cells} ячеек в строке."
        )

    # гарантируем, что строк хватит
    needed_rows = start_row + ((len(values) + cols_per_row - 1) // cols_per_row)
    while len(table.rows) < needed_rows:
        table.add_row()

    for idx, value in enumerate(values):
        r = start_row + (idx // cols_per_row)
        c = start_col + (idx % cols_per_row)
        table.rows[r].cells[c].text = str(value)

    # На Windows docx часто блокируется Word'ом. Если нельзя перезаписать файл —
    # сохраняем рядом под новым именем.
    try:
        doc.save(result_docx_path)
        return result_docx_path
    except PermissionError:
        p = Path(result_docx_path)
        alt = p.with_name(f"{p.stem}_filled{p.suffix}")
        if alt.exists():
            import uuid as _uuid
            alt = p.with_name(f"{p.stem}_filled_{str(_uuid.uuid4())[:8]}{p.suffix}")
        doc.save(str(alt))
        print(
            f"Не удалось перезаписать '{result_docx_path}' (файл, вероятно, открыт). "
            f"Сохранил результат в '{alt.name}'. Закройте Word и запустите снова, если нужно перезаписать исходный файл."
        )
        return str(alt)


def fill_result_table_from_perplexity(
    result_docx_path: str = "result.docx",
    model: str = "sonar",
    thread_id: Optional[str] = "7b33ace1-39ff-4ab1-aa21-29e796bff58a",
    store_path: str = DEFAULT_CHAT_STORE,
    table_start_row: int = 1,
    table_start_col: int = 0,
    cols_per_row: int = 3,
    table_index: Optional[int] = None,
) -> List[str]:
    """
    1) Отправляет 2 файла в Perplexity (как контент) и промпт (пока заглушка).
    2) Ожидает, что модель вернёт список значений.
    3) Заполняет таблицу под индексом 1 в docx-файле result_docx_path.

    Возвращает список значений, которыми заполнили таблицу.
    """

    if not os.path.exists(result_docx_path):
        raise ValueError(f"Файл результата не найден: {result_docx_path}. Сначала сгенерируйте result.docx.")

    PROMPT = (
        "Теперь, когда у тебя есть шаблон РПД и материалы, у меня есть для тебя еще одно задание. В шаблоне РПД есть несколько таблиц. Твоя задача - подобрать из учебных материалов, либо, если там напрямую не указано, проанализировать материалы и предположить данные для заполнения таблиц. Например, в первой присутствуют такие столбцы как: Категория (группа) компетенции, Код и наименование компетенции, Код и наименование индикатора достижения компетенции. Твой ответ с данными (будь то напрямую указанные из материалов или предположенные) должен быть в формате json списка списков, где, например, для первой таблицы в ее списке: 0 - индекс первого столбца, 1 - индекс второго, 2 - третьего, 3 - опять первого столбца, но второй строки, 4 - второго столбца, второй строки и так далее (данных может быть сколько угодно по количеству строк и тд. главное - логичность). Данные по остальным таблицам нужно организовывать в таком же формате в своих собственных списках внутри общего, но с учетом их индивидуального количества столбцов и соответственно своей индексацией. Не нужно нигде писать сам ты предположил данные после анализа или взял напрямую из материалов. Просто сухой четкий ответ по формату"
    )

    # Локальный чат для сохранения контекста (как в request_api.py)
    chat_id = thread_id
    if not chat_id:
        # если не передали — создадим новый
        import uuid as _uuid
        chat_id = str(_uuid.uuid4())
    print(f"CHAT_ID: {chat_id}")

    messages = _load_chat_messages(chat_id, store_path=store_path)
    if not any(m.get("role") == "system" for m in messages):
        messages.insert(0, {"role": "system", "content": "Be precise and concise."})

    messages.append(
        {
            "role": "user",
            "content": (
                f"{PROMPT}\n\n"
            ),
        }
    )

    stream = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.2,
        stream=True,
    )

    completion_id = None
    answer_text = ""
    print("Получение списка значений от API...")
    for chunk in stream:
        if completion_id is None and hasattr(chunk, "id") and chunk.id:
            completion_id = chunk.id
            print(f"COMPLETION_ID: {completion_id}")
        if getattr(chunk, "choices", None) and len(chunk.choices) > 0:
            delta = chunk.choices[0].delta
            content = getattr(delta, "content", None)
            if content:
                answer_text += content
                print(content, end="", flush=True)
    print()

    print("\n--- RAW_TABLE_RESPONSE (first 500 chars) ---")
    print(answer_text[:500])
    print("--- /RAW_TABLE_RESPONSE ---\n")

    values = _extract_list_values(answer_text)
    print(f"Распарсено значений: {len(values)} (cols_per_row={cols_per_row}, start={table_start_row}:{table_start_col})")
    if values:
        print("Первые 10 значений:", values[:10])
    if not values:
        raise ValueError("Модель не вернула список значений (не удалось распарсить ответ).")

    # сохраняем историю
    messages.append({"role": "assistant", "content": answer_text})
    _save_chat_messages(chat_id, messages, store_path=store_path)

    # Определяем, какую таблицу заполнять.
    # По умолчанию пытаемся найти таблицу по заголовкам компетенций.
    if table_index is None:
        table_index = find_table_index_by_headers(
            result_docx_path,
            header_substrings=[
                "Категория",
                "Код и наименование компетенции",
                "Код и наименование индикатора",
            ],
            header_row=0,
        )
        print(f"Найдена таблица по заголовкам. table_index={table_index}")

    # Заполняем выбранную таблицу слева-направо по cols_per_row колонок,
    # начиная с table_start_row:table_start_col (у вас из‑за заголовков это 1:0).
    fill_table_row_major(
        result_docx_path,
        values,
        table_index=table_index,
        cols_per_row=cols_per_row,
        start_row=table_start_row,
        start_col=table_start_col,
    )
    return values

