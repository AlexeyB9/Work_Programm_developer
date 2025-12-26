import json
import os
import uuid
from pathlib import Path

from docx import Document
from openai import OpenAI

# API ключ должен быть установлен через переменную окружения PPLX_API_KEY
PPLX_API_KEY = os.getenv("PPLX_API_KEY")
if not PPLX_API_KEY:
    raise ValueError(
        "PPLX_API_KEY не установлен. "
        "Установите переменную окружения: export PPLX_API_KEY=your_key "
        "или настройте на сервере через панель управления хостинга."
    )

client = OpenAI(
    api_key=PPLX_API_KEY,
    base_url="https://api.perplexity.ai",
)

# Путь к файлу истории чатов в корне проекта (рядом с main.py)
DEFAULT_CHAT_STORE = str(Path(__file__).resolve().parent.parent / "perplexity_chats.json")


def _load_chat_messages(chat_id: str, store_path: str = DEFAULT_CHAT_STORE) -> list[dict]:
    store_file = Path(store_path)
    if not store_file.exists():
        return []

    try:
        data = json.loads(store_file.read_text(encoding="utf-8"))
    except Exception:
        return []

    msgs = data.get(chat_id, [])
    return msgs if isinstance(msgs, list) else []


def _save_chat_messages(chat_id: str, messages: list[dict], store_path: str = DEFAULT_CHAT_STORE) -> None:
    store_file = Path(store_path)
    data: dict = {}

    if store_file.exists():
        try:
            data = json.loads(store_file.read_text(encoding="utf-8"))
        except Exception:
            data = {}

    data[chat_id] = messages
    store_file.parent.mkdir(parents=True, exist_ok=True)
    store_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def read_docx_file(file_path: str) -> str:
    """
    Читает содержимое DOCX или DOTM файла.

    Args:
        file_path: путь к файлу

    Returns:
        текстовое содержимое файла
    """

    try:
        doc = Document(file_path)
        text_parts = []

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts) if text_parts else ""
    except Exception as e:
        raise ValueError(f"Не удалось прочитать DOCX файл {file_path}: {e}")


def read_file_content(file_path: str) -> str:
    """
    Читает содержимое файла с автоматическим определением формата и кодировки.
    Поддерживает текстовые файлы, DOCX и DOTM.

    Args:
        file_path: путь к файлу

    Returns:
        содержимое файла в виде строки

    Raises:
        ValueError: если файл не может быть прочитан
    """
    # Определяем расширение файла
    file_ext = os.path.splitext(file_path)[1].lower()

    # Обработка DOCX и DOTM файлов
    if file_ext in ['.docx', '.dotm']:
        return read_docx_file(file_path)

    # Обработка текстовых файлов
    # Список кодировок для попытки чтения
    encodings = ["utf-8", "cp1251", "latin-1", "iso-8859-1"]

    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            raise ValueError(f"Не удалось прочитать файл {file_path}: {e}")

    # Если все кодировки не подошли, пробуем прочитать как бинарный и предупреждаем
    try:
        with open(file_path, "rb") as f:
            content = f.read()
            # Проверяем, является ли файл текстовым (первые байты)
            if b'\x00' in content[:1024]:  # Наличие нулевых байтов указывает на бинарный файл
                raise ValueError(
                    f"Файл {file_path} является бинарным. "
                    f"Поддерживаются только текстовые файлы и файлы формата DOCX/DOTM."
                )
            # Пробуем декодировать как UTF-8 с обработкой ошибок
            return content.decode("utf-8", errors="replace")
    except Exception as e:
        raise ValueError(f"Не удалось прочитать файл {file_path}: {e}")


def call_api_in_one(
        file1_path: str,
        file2_path: str,
        prompt: str,
        model: str = "sonar",
        thread_id: str | None = None,
        store_path: str = DEFAULT_CHAT_STORE,
) -> tuple[str, str]:
    """
    Отправляет два файла и промпт в API Perplexity с потоковой обработкой и возвращает ответ.

    Важно: Perplexity Chat Completions — OpenAI-compatible, но у него **нет гарантированного server-side thread_id**,
    который можно передать потом и продолжить "тот же чат" как в Assistants API.
    Поэтому для "подключения к тому же чату" мы делаем правильно и надежно:
    - печатаем локальный **CHAT_ID**
    - сохраняем историю `messages` в файл `perplexity_chats.json`
    - при следующем вызове с тем же `thread_id` (CHAT_ID) подхватываем историю и продолжаем контекст

    Args:
        file1_path: путь к первому файлу
        file2_path: путь ко второму файлу
        prompt: промпт для обработки файлов
        model: модель Perplexity (по умолчанию "sonar", также доступна "sonar-pro")
        thread_id: CHAT_ID для продолжения разговора (опционально)
        store_path: путь к файлу-хранилищу истории чатов

    Returns:
        (answer_text, chat_id)
        - answer_text: ответ модели
        - chat_id: локальный CHAT_ID, по которому сохраняется/подхватывается история messages
    """
    chat_id = thread_id or str(uuid.uuid4())
    print(f"CHAT_ID: {chat_id}")

    messages: list[dict] = _load_chat_messages(chat_id, store_path=store_path)
    print(f"Загружено сообщений из истории: {len(messages)} (store: {store_path})")

    if not any(m.get("role") == "system" for m in messages):
        messages.insert(
            0,
            {"role": "system", "content": "Вы — полезный ассистент, который анализирует файлы и отвечает на вопросы."},
        )

    # Читаем содержимое файлов
    file1_content = read_file_content(file1_path)
    file2_content = read_file_content(file2_path)

    messages.append(
        {
            "role": "user",
            "content": (
                f"Файл 1 ({os.path.basename(file1_path)}):\n{file1_content}\n\n"
                f"Файл 2 ({os.path.basename(file2_path)}):\n{file2_content}\n\n"
                f"Промпт: {prompt}"
            ),
        }
    )

    # Параметры для потокового запроса
    stream_params = {
        "model": model,
        "messages": messages,
        "temperature": 0.4,
        "stream": True
    }

    # Отправляем потоковый запрос с обработкой ошибок
    try:
        stream = client.chat.completions.create(**stream_params)
    except Exception as api_error:
        error_msg = str(api_error)
        if "api_key" in error_msg.lower() or "authentication" in error_msg.lower() or "401" in error_msg:
            raise ValueError(
                f"Ошибка аутентификации API: Проверьте что PPLX_API_KEY установлен правильно. "
                f"Детали: {error_msg}"
            )
        elif "connection" in error_msg.lower() or "timeout" in error_msg.lower() or "network" in error_msg.lower():
            raise ConnectionError(
                f"Ошибка подключения к Perplexity API: {error_msg}. "
                f"Проверьте интернет-соединение и доступность api.perplexity.ai"
            )
        else:
            raise Exception(f"Ошибка при запросе к Perplexity API: {error_msg}")

    # В streaming у чанков обычно есть chunk.id (completion id). Это НЕ chat/thread id, но полезно для логов.
    completion_id: str | None = None
    full_response = ""

    for chunk in stream:
        if completion_id is None and hasattr(chunk, "id") and chunk.id:
            completion_id = chunk.id

        # Собираем содержимое ответа
        if getattr(chunk, "choices", None) and len(chunk.choices) > 0:
            delta = chunk.choices[0].delta
            content = None

            if hasattr(delta, "content"):
                content = delta.content
            elif isinstance(delta, dict):
                content = delta.get("content")

            if content:
                full_response += content

    # Сохраняем историю для продолжения "того же чата" через CHAT_ID
    if full_response:
        messages.append({"role": "assistant", "content": full_response})
        _save_chat_messages(chat_id, messages, store_path=store_path)

    return (full_response if full_response else "Ответ не содержит данных.", chat_id)


def call_api_in_two(
        file1_path: str,
        prompt: str,
        model: str = "sonar",
        thread_id: str | None = None,
        store_path: str = DEFAULT_CHAT_STORE,
) -> str:
    """
    Отправляет один файл и промпт в API Perplexity с потоковой обработкой и возвращает ответ.

    Важно: Perplexity Chat Completions — OpenAI-compatible, но у него **нет гарантированного server-side thread_id**,
    который можно передать потом и продолжить "тот же чат" как в Assistants API.
    Поэтому для "подключения к тому же чату" мы делаем правильно и надежно:
    - печатаем локальный **CHAT_ID**
    - сохраняем историю `messages` в файл `perplexity_chats.json`
    - при следующем вызове с тем же `thread_id` (CHAT_ID) подхватываем историю и продолжаем контекст

    Args:
        file1_path: путь к первому файлу
        prompt: промпт для обработки файлов
        model: модель Perplexity (по умолчанию "sonar", также доступна "sonar-pro")
        thread_id: CHAT_ID для продолжения разговора (опционально)
        store_path: путь к файлу-хранилищу истории чатов

    Returns:
        ответ от API в виде строки
    """

    if not thread_id:
        raise ValueError(
            "Для call_api_in_two нужно передать thread_id (это ваш CHAT_ID из предыдущего запуска)."
        )

    messages: list[dict] = _load_chat_messages(thread_id, store_path=store_path)
    if not messages:
        raise ValueError(
            "История чата не найдена для указанного CHAT_ID.\n"
            f"- CHAT_ID: {thread_id}\n"
            f"- store: {store_path}\n"
            "Сначала запустите call_api_in_one и возьмите CHAT_ID из консоли (и убедитесь, что сохранился perplexity_chats.json)."
        )

    print(f"CHAT_ID: {thread_id}")
    print(f"Загружено сообщений из истории: {len(messages)} (store: {store_path})")

    if not any(m.get("role") == "system" for m in messages):
        messages.insert(
            0,
            {"role": "system", "content": "Вы — полезный ассистент, который анализирует файлы и отвечает на вопросы."},
        )

    # Читаем содержимое файлов
    file1_content = read_file_content(file1_path)


    messages.append(
        {
            "role": "user",
            "content": (
                f"Файл 1 ({os.path.basename(file1_path)}):\n{file1_content}\n\n"
                f"Промпт: {prompt}"
            ),
        }
    )

    # Параметры для потокового запроса
    stream_params = {
        "model": model,
        "messages": messages,
        "temperature": 0.4,
        "stream": True
    }

    # Отправляем потоковый запрос с обработкой ошибок
    try:
        stream = client.chat.completions.create(**stream_params)
    except Exception as api_error:
        error_msg = str(api_error)
        if "api_key" in error_msg.lower() or "authentication" in error_msg.lower() or "401" in error_msg:
            raise ValueError(
                f"Ошибка аутентификации API: Проверьте что PPLX_API_KEY установлен правильно. "
                f"Детали: {error_msg}"
            )
        elif "connection" in error_msg.lower() or "timeout" in error_msg.lower() or "network" in error_msg.lower():
            raise ConnectionError(
                f"Ошибка подключения к Perplexity API: {error_msg}. "
                f"Проверьте интернет-соединение и доступность api.perplexity.ai"
            )
        else:
            raise Exception(f"Ошибка при запросе к Perplexity API: {error_msg}")

    # В streaming у чанков обычно есть chunk.id (completion id). Это НЕ chat/thread id, но полезно для логов.
    completion_id: str | None = None
    full_response = ""

    for chunk in stream:
        if completion_id is None and hasattr(chunk, "id") and chunk.id:
            completion_id = chunk.id

        # Собираем содержимое ответа
        if getattr(chunk, "choices", None) and len(chunk.choices) > 0:
            delta = chunk.choices[0].delta
            content = None

            if hasattr(delta, "content"):
                content = delta.content
            elif isinstance(delta, dict):
                content = delta.get("content")

            if content:
                full_response += content

    # Сохраняем историю для продолжения "того же чата" через CHAT_ID
    if full_response:
        messages.append({"role": "assistant", "content": full_response})
        _save_chat_messages(thread_id, messages, store_path=store_path)

    return full_response if full_response else "Ответ не содержит данных."

