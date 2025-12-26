from __future__ import annotations

import json
import os
import re
from typing import List, Optional, Sequence

from docx import Document

from wpd.request_api import DEFAULT_CHAT_STORE, _load_chat_messages, _save_chat_messages, client
from wpd.fill_result_table import fill_table_row_major


def _to_zero_based_table_index(table_index: int, index_base: int, table_index_offset: int = 0) -> int:
    """
    Преобразует table_index из конфигурации в индекс для python-docx (0-based).

    По вашей логике:
    - table_index начинается с 1 (первая таблица = 1)
    В python-docx:
    - первая таблица = doc.tables[0]
    """
    if index_base not in (0, 1):
        raise ValueError("index_base должен быть 0 или 1")
    if index_base == 0:
        base_idx = table_index
    else:
        base_idx = table_index - 1
    return base_idx + table_index_offset


def _stringify_content(content) -> str:
    if content is None:
        return ""
    if isinstance(content, str):
        return content
    try:
        return json.dumps(content, ensure_ascii=False)
    except Exception:
        return str(content)


def _normalize_messages(messages: list[dict]) -> list[dict]:
    """
    Perplexity Chat Completions требует строгую чередуемость ролей:
    после (опциональных) system сообщений дальше должны чередоваться user/tool и assistant.

    Мы делаем упрощённо и совместимо с Perplexity:
    - оставляем только роли system/user/assistant
    - все system поднимаем в начало
    - подряд идущие одинаковые роли (user-user, assistant-assistant) склеиваем в одно сообщение
    """
    if not messages:
        return []

    systems: list[dict] = []
    rest: list[dict] = []

    for m in messages:
        role = (m.get("role") or "").strip()
        if role not in ("system", "user", "assistant"):
            continue
        content = _stringify_content(m.get("content"))
        item = {"role": role, "content": content}
        if role == "system":
            systems.append(item)
        else:
            rest.append(item)

    # Склеиваем подряд одинаковые роли
    merged: list[dict] = []
    for m in rest:
        if not merged:
            merged.append(m)
            continue
        if merged[-1]["role"] == m["role"]:
            prev = merged[-1]["content"].strip()
            cur = m["content"].strip()
            merged[-1]["content"] = f"{prev}\n\n{cur}" if prev and cur else (prev or cur)
        else:
            merged.append(m)

    # Если после system вообще нет user/assistant — вернем только system
    return systems + merged


def _extract_values_from_ai_response(text: str) -> List[str]:
    """
    Универсальный парсер для "значений таблицы" из ответа ИИ.

    Ожидаемый лучший формат: JSON-массив:
      ["a", "b", "c", ...]

    Также поддерживает:
    - JSON list-of-lists -> будет "сплющен" в плоский список
    - маркированные/нумерованные списки
    - строка с ';' или ','
    
    ВАЖНО: Пустые значения сохраняются и не фильтруются.
    """
    if not text:
        return []

    s = text.strip()

    # 1) JSON целиком
    try:
        obj = json.loads(s)
        if isinstance(obj, list):
            # list[str] или list[list[str]] -> flatten
            flat: list[str] = []
            for x in obj:
                if isinstance(x, list):
                    # Сохраняем все значения, включая пустые
                    flat.extend(str(v).strip() if v is not None else "" for v in x)
                else:
                    # Сохраняем значение, включая пустое
                    flat.append(str(x).strip() if x is not None else "")
            return flat  # Возвращаем все значения, включая пустые
    except Exception:
        pass

    # 2) JSON как подстрока
    m = re.search(r"\[[\s\S]*\]", s)
    if m:
        try:
            obj = json.loads(m.group(0))
            if isinstance(obj, list):
                flat: list[str] = []
                for x in obj:
                    if isinstance(x, list):
                        # Сохраняем все значения, включая пустые
                        flat.extend(str(v).strip() if v is not None else "" for v in x)
                    else:
                        # Сохраняем значение, включая пустое
                        flat.append(str(x).strip() if x is not None else "")
                return flat  # Возвращаем все значения, включая пустые
        except Exception:
            pass

    # 3) списки по строкам
    out: list[str] = []
    for line in s.replace("\r\n", "\n").split("\n"):
        line = line.strip()
        # Пропускаем только полностью пустые строки между элементами
        # Но если это единственный элемент, сохраняем его
        if not line:
            continue
        line = re.sub(r"^[-•]\s*", "", line)
        line = re.sub(r"^\d+[\.\)]\s*", "", line)
        # Сохраняем строку, даже если она стала пустой после удаления маркеров
        out.append(line)
    if len(out) >= 2:
        return out

    # 4) разделители
    if ";" in s:
        # Сохраняем все значения, включая пустые между разделителями
        return [p.strip() for p in s.split(";")]
    if "," in s:
        # Сохраняем все значения, включая пустые между разделителями
        return [p.strip() for p in s.split(",")]

    return [s] if s else []


def fill_one_table_from_perplexity(
    *,
    result_docx_path: str,
    table_index: int,
    cols_per_row: int,
    start_row: int,
    start_col: int,
    prompt: str,
    model: str = "sonar",
    thread_id: Optional[str] = None,
    store_path: str = DEFAULT_CHAT_STORE,
    index_base: int = 1,
    table_index_offset: int = 0,
) -> List[str]:
    """
    Универсальная функция для заполнения ОДНОЙ таблицы.

    Принимает:
    - table_index: индекс таблицы в docx (doc.tables[table_index])
    - cols_per_row: количество колонок (ширина области заполнения)
    - start_row/start_col: стартовая точка (например 1:0 если 0-я строка — заголовки)
    - prompt: промпт (передаётся снаружи)

    Важно: Perplexity Chat Completions контекст держит только через `messages`,
    поэтому мы используем `thread_id` как локальный CHAT_ID для загрузки истории.
    """
    if not os.path.exists(result_docx_path):
        raise ValueError(f"Файл результата не найден: {result_docx_path}")
    if not thread_id:
        raise ValueError("Нужно передать thread_id (CHAT_ID), чтобы модель имела контекст предыдущих файлов/сообщений.")

    messages = _load_chat_messages(thread_id, store_path=store_path)
    if not messages:
        raise ValueError(
            f"История чата пуста для CHAT_ID={thread_id}. "
            "Сначала запустите шаг, который загружает шаблон/материалы в чат."
        )

    # Нормализуем историю, чтобы Perplexity не ругался на порядок ролей (400 invalid_message)
    messages = _normalize_messages(messages)

    # Добавляем новый промпт: если последний user — дописываем, иначе добавляем новым user
    if messages and messages[-1]["role"] == "user":
        prev = messages[-1]["content"].strip()
        messages[-1]["content"] = f"{prev}\n\n{prompt}" if prev else prompt
    else:
        messages.append({"role": "user", "content": prompt})

    # Минимальное логирование для оптимизации
    print(f"\n=== TABLE {table_index} (start {start_row}:{start_col}, cols {cols_per_row}) ===")
    print(f"История чата: {len(messages)} сообщений")
    
    # Отправляем запрос точно так же, как в call_api_in_one
    stream_params = {
        "model": model,
        "messages": messages,
        "temperature": 0.2,
        "stream": True
    }
    
    try:
        stream = client.chat.completions.create(**stream_params)
    except Exception as api_error:
        error_msg = str(api_error)
        print(f"ОШИБКА при запросе к API для таблицы {table_index}: {error_msg}")
        import traceback
        traceback.print_exc()
        if "api_key" in error_msg.lower() or "authentication" in error_msg.lower() or "401" in error_msg:
            raise ValueError(
                f"Ошибка аутентификации API для таблицы {table_index}: Проверьте что PPLX_API_KEY установлен правильно. "
                f"Детали: {error_msg}"
            )
        elif "connection" in error_msg.lower() or "timeout" in error_msg.lower() or "network" in error_msg.lower():
            raise ConnectionError(
                f"Ошибка подключения к Perplexity API для таблицы {table_index}: {error_msg}. "
                f"Проверьте интернет-соединение и доступность api.perplexity.ai"
            )
        else:
            raise Exception(f"Ошибка при запросе к API для таблицы {table_index}: {error_msg}")

    completion_id = None
    answer_text = ""
    
    try:
        for chunk in stream:
            if completion_id is None and hasattr(chunk, "id") and chunk.id:
                completion_id = chunk.id
            
            if getattr(chunk, "choices", None) and len(chunk.choices) > 0:
                delta = chunk.choices[0].delta
                content = getattr(delta, "content", None)
                if content:
                    answer_text += content
    except Exception as stream_error:
        error_msg = str(stream_error)
        if not answer_text:
            raise Exception(f"Не удалось получить ответ от API для таблицы {table_index}: {error_msg}")

    # Преобразуем номер таблицы из вашей схемы (1-based) в python-docx индекс (0-based)
    doc_table_index = _to_zero_based_table_index(
        table_index,
        index_base=index_base,
        table_index_offset=table_index_offset,
    )

    # Диагностика: какая таблица реально будет изменена и сколько колонок она имеет по мнению python-docx
    try:
        d = Document(result_docx_path)
        if doc_table_index < 0 or doc_table_index >= len(d.tables):
            raise ValueError(
                f"table_index={table_index} (index_base={index_base}) -> doc_index={doc_table_index} вне диапазона. "
                f"Всего таблиц в документе: {len(d.tables)}"
            )
        t = d.tables[doc_table_index]
        max_cells = max((len(r.cells) for r in t.rows), default=0)
        header = " | ".join(c.text.strip() for c in (t.rows[0].cells if t.rows else []))[:120]
        print(
            f"[TABLE] table_index={table_index} (base={index_base}, offset={table_index_offset}) -> doc_index={doc_table_index}; "
            f"rows={len(t.rows)} len(columns)={len(t.columns)} max_row_cells={max_cells}; header='{header}'"
        )
    except Exception as e:
        print(f"[TABLE] Не удалось прочитать структуру таблицы перед заполнением: {e}")

    # Парсим значения из ответа ИИ
    values = _extract_values_from_ai_response(answer_text)
    if not values:
        raise ValueError(f"Не удалось распарсить значения из ответа ИИ для таблицы {table_index} (ожидался JSON-массив). Ответ был: {answer_text[:200]}...")

    # Сохраняем в историю (и сразу нормализуем, чтобы не копить "ломаную" последовательность)
    messages.append({"role": "assistant", "content": answer_text})
    messages = _normalize_messages(messages)
    _save_chat_messages(thread_id, messages, store_path=store_path)

    # заполняем таблицу (ВАЖНО: table_index здесь уже doc_index)
    fill_table_row_major(
        result_docx_path=result_docx_path,
        values=values,
        table_index=doc_table_index,
        cols_per_row=cols_per_row,
        start_row=start_row,
        start_col=start_col,
    )
    return values


def fill_tables_from_lists(
    *,
    result_docx_path: str,
    table_indices: Sequence[int],
    cols_per_row_list: Sequence[int],
    start_coords: Sequence[tuple[int, int]],
    prompts: Sequence[str],
    model: str = "sonar",
    thread_id: Optional[str],
    store_path: str = DEFAULT_CHAT_STORE,
    index_base: int = 1,
    table_index_offset: int = 0,
) -> None:
    """
    Вызов в цикле по параллельным спискам (как вы описали).
    """
    if not (len(table_indices) == len(cols_per_row_list) == len(start_coords) == len(prompts)):
        raise ValueError("Длины списков table_indices / cols_per_row_list / start_coords / prompts должны совпадать.")

    for i in range(len(table_indices)):
        r, c = start_coords[i]
        fill_one_table_from_perplexity(
            result_docx_path=result_docx_path,
            table_index=table_indices[i],
            cols_per_row=cols_per_row_list[i],
            start_row=r,
            start_col=c,
            prompt=prompts[i],
            model=model,
            thread_id=thread_id,
            store_path=store_path,
            index_base=index_base,
            table_index_offset=table_index_offset,
        )


